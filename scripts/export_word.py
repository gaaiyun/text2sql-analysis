"""
Word 导出工具 - 场景 5 企业尽调报告

使用方法:
    python export_word.py report_data.json output.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import sys

def create_due_diligence_report(data, filename, template_path=None):
    """
    创建企业尽调报告（场景 5）
    
    参数:
        data: dict 包含企业信息和分析结果
        filename: 输出文件名
        template_path: 可选的模板文件路径
    
    报告结构:
    1. 封面
    2. 企业基本信息
    3. 经营状况分析
    4. 知识产权情况
    5. 法律风险分析
    6. 投融资情况
    7. 综合评估与建议
    """
    
    if template_path:
        doc = Document(template_path)
    else:
        doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    # 1. 封面
    title = doc.add_heading('企业尽职调查报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"企业名称：{data.get('企业名称', '未知')}")
    run.font.size = Pt(16)
    run.font.bold = True
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"报告日期：{data.get('报告日期', '')}")
    date_run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # 2. 企业基本信息
    doc.add_heading('一、企业基本信息', level=1)
    
    basic_info = data.get('基本信息', {})
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '项目'
    hdr_cells[1].text = '内容'
    
    for key, value in basic_info.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value)
    
    # 3. 经营状况分析
    doc.add_heading('二、经营状况分析', level=1)
    
    operation = data.get('经营状况', {})
    for key, value in operation.items():
        p = doc.add_paragraph()
        p.add_run(f'{key}：').bold = True
        p.add_run(str(value))
    
    # 4. 知识产权情况
    doc.add_heading('三、知识产权情况', level=1)
    
    ip = data.get('知识产权', {})
    ip_table = doc.add_table(rows=1, cols=2)
    ip_table.style = 'Light Grid Accent 1'
    
    ip_hdr = ip_table.rows[0].cells
    ip_hdr[0].text = '类型'
    ip_hdr[1].text = '数量'
    
    for key, value in ip.items():
        row = ip_table.add_row().cells
        row[0].text = key
        row[1].text = str(value)
    
    # 5. 法律风险分析
    doc.add_heading('四、法律风险分析', level=1)
    
    legal = data.get('法律风险', {})
    if legal.get('诉讼数量', 0) > 0:
        doc.add_paragraph(f"⚠️ 存在 {legal.get('诉讼数量')} 起诉讼案件", style='List Bullet')
        for case in legal.get('诉讼详情', []):
            doc.add_paragraph(f"- {case}", style='List Bullet 2')
    else:
        doc.add_paragraph("✅ 无重大法律诉讼记录", style='List Bullet')
    
    # 6. 投融资情况
    doc.add_heading('五、投融资情况', level=1)
    
    financing = data.get('投融资', [])
    if financing:
        fin_table = doc.add_table(rows=1, cols=4)
        fin_table.style = 'Light Grid Accent 1'
        
        fin_hdr = fin_table.rows[0].cells
        fin_hdr[0].text = '日期'
        fin_hdr[1].text = '轮次'
        fin_hdr[2].text = '金额'
        fin_hdr[3].text = '投资方'
        
        for item in financing:
            row = fin_table.add_row().cells
            row[0].text = item.get('日期', '')
            row[1].text = item.get('轮次', '')
            row[2].text = item.get('金额', '')
            row[3].text = item.get('投资方', '')
    else:
        doc.add_paragraph("暂无公开融资记录")
    
    # 7. 综合评估与建议
    doc.add_heading('六、综合评估与建议', level=1)
    
    score = data.get('综合评分', 0)
    rating = '优秀' if score >= 80 else '良好' if score >= 60 else '一般'
    
    p = doc.add_paragraph()
    p.add_run(f'综合评分：{score} 分 ({rating})\n\n').bold = True
    
    suggestions = data.get('建议', [])
    doc.add_paragraph('投资建议：')
    for suggestion in suggestions:
        doc.add_paragraph(suggestion, style='List Number')
    
    # 保存文档
    doc.save(filename)
    print(f"✅ 生成报告: {filename}")
    return filename

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python export_word.py report_data.json output.docx")
        sys.exit(1)
    
    with open(sys.argv[1], 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    create_due_diligence_report(data, sys.argv[2])
