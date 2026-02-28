#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
文档生成模块

支持 PDF（reportlab）、Word（python-docx）、HTML（Jinja2 模板）。
支持嵌入图片和表格，供场景 1-5 多模态输出使用。
"""

import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime

# PDF
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    Image, PageBreak,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Word
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# HTML
from jinja2 import Environment, BaseLoader, select_autoescape


def _ensure_output_dir(output_path: str) -> str:
    p = Path(output_path)
    p.parent.mkdir(parents=True, exist_ok=True)
    return str(p.resolve())


# ---------- PDF ----------

def _register_chinese_font():
    """注册中文字体（若系统有），否则用 Helvetica 避免报错"""
    try:
        simhei = Path("C:/Windows/Fonts/simhei.ttf")
        if simhei.exists():
            pdfmetrics.registerFont(TTFont("SimHei", str(simhei)))
            return "SimHei"
    except Exception:
        pass
    return "Helvetica"


def pdf_from_markdown_sections(
    sections: List[Dict[str, Any]],
    output_path: str,
    title: str = "分析报告",
    image_paths: Optional[List[str]] = None,
    page_size=A4,
) -> str:
    """
    根据「标题 + 内容」段落和可选图片列表生成 PDF。

    Args:
        sections: [{"heading": "一级标题", "content": "段落或 markdown 文本"}, ...]
        output_path: 输出 PDF 路径
        title: 文档总标题
        image_paths: 要按顺序插入的图片路径列表

    Returns:
        生成的 PDF 文件路径
    """
    path = _ensure_output_dir(output_path)
    doc = SimpleDocTemplate(
        path,
        pagesize=page_size,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm,
    )
    font_name = _register_chinese_font()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        name="CustomTitle",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=18,
        spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        name="CustomHeading",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=14,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        name="CustomBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=10,
        spaceAfter=6,
    )
    story = []
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 12))

    image_index = 0
    for sec in sections:
        if sec.get("heading"):
            story.append(Paragraph(sec["heading"], heading_style))
        content = sec.get("content", "")
        if content:
            # 简单按行处理，避免 reportlab 不支持的 markdown 语法
            for line in content.strip().split("\n"):
                line = line.strip()
                if not line:
                    continue
                story.append(Paragraph(line.replace("<", "&lt;").replace(">", "&gt;"), body_style))
        # 若该 section 指定了图片或使用全局 image_paths
        sec_images = sec.get("images", [])
        if not sec_images and image_paths and image_index < len(image_paths):
            sec_images = image_paths[image_index:image_index+1]
            image_index += 1
        for img_path in sec_images:
            if os.path.isfile(img_path):
                try:
                    im = Image(img_path, width=14*cm, height=10*cm)
                    story.append(Spacer(1, 6))
                    story.append(im)
                    story.append(Spacer(1, 12))
                except Exception:
                    pass

    if image_paths:
        for i in range(image_index, len(image_paths)):
            if os.path.isfile(image_paths[i]):
                try:
                    story.append(Spacer(1, 6))
                    story.append(Image(image_paths[i], width=14*cm, height=10*cm))
                    story.append(Spacer(1, 12))
                except Exception:
                    pass

    doc.build(story)
    return path


def pdf_from_tables(
    tables: List[Dict[str, Any]],
    output_path: str,
    title: str = "数据报告",
    page_size=A4,
) -> str:
    """
    用表格数据生成 PDF。每个 table: {"headers": [...], "rows": [[...], ...]}。

    Returns:
        PDF 文件路径
    """
    path = _ensure_output_dir(output_path)
    doc = SimpleDocTemplate(
        path,
        pagesize=page_size,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm,
    )
    font_name = _register_chinese_font()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        name="PdfTitle",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=16,
        spaceAfter=12,
    )
    story = [Paragraph(title, title_style), Spacer(1, 12)]

    for t in tables:
        headers = t.get("headers", [])
        rows = t.get("rows", [])
        data = [headers] + rows
        if not data:
            continue
        tbl = Table(data)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
            ("FONTNAME", (0, 0), (-1, -1), font_name),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 16))
    doc.build(story)
    return path


# ---------- Word ----------

def word_from_sections(
    sections: List[Dict[str, Any]],
    output_path: str,
    title: str = "分析报告",
    image_paths: Optional[List[str]] = None,
) -> str:
    """
    用「标题+内容」段落和可选图片生成 Word。

    Args:
        sections: [{"heading": "...", "content": "...", "images": [...]}, ...]
        image_paths: 全局图片列表，按顺序插入

    Returns:
        docx 文件路径
    """
    path = _ensure_output_dir(output_path)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)

    doc.add_heading(title, level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    image_index = 0
    for sec in sections:
        if sec.get("heading"):
            doc.add_heading(sec["heading"], level=1)
        
        content = (sec.get("content") or "").strip()
        if content:
            # 处理markdown格式
            lines = content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # 处理标题
                if line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                # 处理代码块
                elif line.startswith("```"):
                    continue
                # 处理列表
                elif line.startswith("- ") or line.startswith("* "):
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                elif line[0:2].isdigit() and line[2:4] in [". ", ") "]:
                    p = doc.add_paragraph(line[3:], style='List Number')
                # 普通段落
                else:
                    # 移除markdown加粗标记
                    line = line.replace("**", "")
                    doc.add_paragraph(line)
        
        sec_images = sec.get("images", [])
        if not sec_images and image_paths and image_index < len(image_paths):
            sec_images = image_paths[image_index:image_index+1]
            image_index += 1
        for img_path in sec_images:
            if os.path.isfile(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(5.5))
                    doc.add_paragraph()
                except Exception:
                    pass

    if image_paths:
        for i in range(image_index, len(image_paths)):
            if os.path.isfile(image_paths[i]):
                try:
                    doc.add_picture(image_paths[i], width=Inches(5.5))
                    doc.add_paragraph()
                except Exception:
                    pass
    doc.save(path)
    return path


def word_add_table(doc: Document, headers: List[str], rows: List[List[Any]]) -> None:
    """向已打开的 Document 添加一张表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
    for row in rows:
        r = table.add_row().cells
        for i, v in enumerate(row):
            if i < len(r):
                r[i].text = str(v)


# ---------- HTML ----------

DEFAULT_HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ title }}</title>
    <style>
        body { font-family: "Microsoft YaHei", sans-serif; margin: 2em; max-width: 900px; }
        h1 { color: #1a237e; border-bottom: 2px solid #1a237e; padding-bottom: 0.3em; }
        h2 { color: #283593; margin-top: 1.5em; }
        table { border-collapse: collapse; width: 100%; margin: 1em 0; }
        th, td { border: 1px solid #ddd; padding: 8px 12px; text-align: left; }
        th { background: #e8eaf6; }
        tr:nth-child(even) { background: #f5f5f5; }
        img { max-width: 100%; height: auto; margin: 1em 0; }
        .section { margin-bottom: 2em; }
        .meta { color: #666; font-size: 0.9em; }
    </style>
</head>
<body>
    <h1>{{ title }}</h1>
    <p class="meta">生成时间：{{ generated_at }}</p>

    {% for section in sections %}
    <div class="section">
        {% if section.heading %}
        <h2>{{ section.heading }}</h2>
        {% endif %}
        {% if section.content %}
        <div>{{ section.content | safe }}</div>
        {% endif %}
        {% if section.table %}
        <table>
            <thead><tr>{% for h in section.table.headers %}<th>{{ h }}</th>{% endfor %}</tr></thead>
            <tbody>
            {% for row in section.table.rows %}
            <tr>{% for c in row %}<td>{{ c }}</td>{% endfor %}</tr>
            {% endfor %}
            </tbody>
        </table>
        {% endif %}
        {% if section.images %}
        {% for img in section.images %}
        <img src="{{ img }}" alt="图表" />
        {% endfor %}
        {% endif %}
    </div>
    {% endfor %}
</body>
</html>
"""


def html_from_sections(
    sections: List[Dict[str, Any]],
    output_path: str,
    title: str = "分析报告",
    template: Optional[str] = None,
) -> str:
    """
    用 sections 生成 HTML。每个 section 可有 heading、content（可 HTML）、table、images。

    Args:
        sections: [{"heading": "...", "content": "...", "table": {"headers": [], "rows": []}, "images": ["path or url"]}, ...]
        template: Jinja2 模板字符串，默认使用内置模板

    Returns:
        HTML 文件路径
    """
    path = _ensure_output_dir(output_path)
    env = Environment(loader=BaseLoader(), autoescape=select_autoescape())
    tpl = env.from_string(template or DEFAULT_HTML_TEMPLATE)
    # 图片转为可访问路径：若为本地路径，转为 file:// 或相对路径
    out_dir = str(Path(path).parent)
    for sec in sections:
        imgs = sec.get("images", [])
        resolved = []
        for img in imgs:
            if os.path.isabs(img) and os.path.isfile(img):
                resolved.append(img)
            elif os.path.isfile(os.path.join(out_dir, img)):
                resolved.append(img)
            else:
                resolved.append(img)
        sec["images"] = resolved
    html = tpl.render(
        title=title,
        sections=sections,
        generated_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    )
    with open(path, "w", encoding="utf-8") as f:
        f.write(html)
    return path
